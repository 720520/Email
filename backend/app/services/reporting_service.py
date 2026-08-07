"""报表字段归集、合同提取、净值指标计算和输入快照构建。"""

from __future__ import annotations

import hashlib
import io
import math
import re
from collections import defaultdict
from datetime import date, timedelta
from decimal import Decimal
from pathlib import Path
from statistics import mean, pstdev
from typing import Any

from docx import Document
from pypdf import PdfReader
from sqlalchemy import exists, or_, select
from sqlalchemy.orm import Session

from app.core.errors import AppError
from app.db.models import FundNav, FundProduct

DEFAULT_REPORT_SECTIONS = [
    "product_info",
    "performance",
    "nav_chart",
    "strategy",
    "contract_terms",
    "disclaimer",
]

REPORT_FIELD_DEFINITIONS: dict[str, dict[str, Any]] = {
    "product_name": {"label": "产品名称", "editable": False, "group": "基础信息"},
    "product_code": {"label": "产品代码/备案代码", "editable": False, "group": "基础信息"},
    "inception_date": {"label": "成立日期", "editable": True, "group": "基础信息"},
    "strategy_category": {"label": "策略分类", "editable": True, "group": "基础信息"},
    "investment_manager": {"label": "投资经理", "editable": True, "group": "策略信息"},
    "investment_strategy": {"label": "策略介绍", "editable": True, "group": "策略信息"},
    "slogan": {"label": "产品摘要", "editable": True, "group": "策略信息"},
    "manager_name": {"label": "管理机构", "editable": True, "group": "合同要素"},
    "custodian_name": {"label": "托管/外包机构", "editable": True, "group": "合同要素"},
    "risk_level": {"label": "风险等级", "editable": True, "group": "合同要素"},
    "open_day": {"label": "开放日", "editable": True, "group": "合同要素"},
    "duration": {"label": "存续期间", "editable": True, "group": "合同要素"},
    "lockup_period": {"label": "锁定期", "editable": True, "group": "合同要素"},
    "management_fee": {"label": "管理费", "editable": True, "group": "费率"},
    "custody_fee": {"label": "托管/外包费", "editable": True, "group": "费率"},
    "subscription_fee": {"label": "申购费", "editable": True, "group": "费率"},
    "redemption_fee": {"label": "赎回费", "editable": True, "group": "费率"},
    "performance_fee": {"label": "业绩报酬", "editable": True, "group": "费率"},
    "investment_scope": {"label": "投资范围", "editable": True, "group": "合同要素"},
    "disclaimer": {"label": "免责声明", "editable": True, "group": "披露信息"},
}

_CONTRACT_PATTERNS: dict[str, tuple[str, ...]] = {
    "product_name": (r"(?:基金|产品)名称\s*[：:]\s*([^\n]{4,100})",),
    "product_code": (r"(?:备案|产品|基金)编码\s*[：:]\s*([A-Za-z0-9_-]{4,64})",),
    "inception_date": (
        r"(?:基金)?成立日期\s*[：:]\s*(\d{4}[年./-]\d{1,2}[月./-]\d{1,2}日?)",
    ),
    "strategy_category": (r"(?:投资)?策略(?:分类|类型)\s*[：:]\s*([^\n；;]{2,80})",),
    "investment_manager": (r"投资经理\s*[：:]\s*([^\n；;]{2,100})",),
    "manager_name": (r"(?:基金)?管理人(?:名称)?\s*[：:]\s*([^\n；;]{4,120})",),
    "custodian_name": (r"(?:托管人|托管机构|外包机构)(?:名称)?\s*[：:]\s*([^\n；;]{2,120})",),
    "risk_level": (r"风险(?:等级|评级)\s*[：:]\s*(R?[1-5]|中高风险|高风险|中风险)",),
    "open_day": (r"开放日\s*[：:]\s*([^\n；;]{2,120})",),
    "duration": (r"(?:基金)?存续(?:期限|期间)\s*[：:]\s*([^\n；;]{2,80})",),
    "lockup_period": (r"锁定期\s*[：:]\s*([^\n；;]{2,80})",),
    "management_fee": (r"管理费(?:率)?\s*[：:]?\s*([0-9.]+%[^\n；;]{0,20})",),
    "custody_fee": (r"(?:托管|外包)(?:服务)?费(?:率)?\s*[：:]?\s*([0-9.]+%[^\n；;]{0,20})",),
    "subscription_fee": (r"(?:申购|认购)费(?:率)?\s*[：:]?\s*([0-9.]+%[^\n；;]{0,50})",),
    "redemption_fee": (r"赎回费(?:率)?\s*[：:]?\s*([0-9.]+%[^\n；;]{0,80})",),
    "performance_fee": (r"(?:业绩报酬|业绩提成)(?:比例)?\s*[：:]?\s*([0-9.]+%[^\n；;]{0,50})",),
    "investment_scope": (r"投资范围\s*[：:]\s*([^\n]{4,500})",),
    "investment_strategy": (r"投资策略\s*[：:]\s*([^\n]{10,1500})",),
}


def extract_contract_text(filename: str, content: bytes) -> str:
    """从 PDF、DOCX 或 TXT 合同中读取文本；扫描件需先做 OCR。"""

    suffix = Path(filename).suffix.casefold()
    try:
        if suffix == ".pdf":
            reader = PdfReader(io.BytesIO(content))
            text = "\n".join((page.extract_text() or "") for page in reader.pages[:500])
        elif suffix == ".docx":
            document = Document(io.BytesIO(content))
            blocks = [paragraph.text for paragraph in document.paragraphs]
            for table in document.tables:
                blocks.extend("\t".join(cell.text for cell in row.cells) for row in table.rows)
            text = "\n".join(blocks)
        elif suffix in {".txt", ".text"}:
            text = _decode_text(content)
        else:
            raise AppError(
                "CONTRACT_FORMAT_UNSUPPORTED",
                "合同目前支持 PDF、DOCX 和 TXT；扫描版 PDF 请先完成 OCR",
            )
    except AppError:
        raise
    except Exception as exc:
        raise AppError("CONTRACT_READ_FAILED", "合同文本读取失败，请检查文件是否损坏") from exc
    normalized = _normalize_contract_text(text)
    if len(normalized) < 20:
        raise AppError(
            "CONTRACT_TEXT_EMPTY",
            "没有从合同中读取到足够文本；扫描版 PDF 请先完成 OCR",
        )
    return normalized[:2_000_000]


def extract_contract_fields(text: str) -> dict[str, str]:
    """按保守规则提取产品要素；未高置信匹配的字段留给人工复核。"""

    result: dict[str, str] = {}
    for key, patterns in _CONTRACT_PATTERNS.items():
        for pattern in patterns:
            match = re.search(pattern, text, flags=re.IGNORECASE)
            if match:
                value = re.sub(r"\s+", " ", match.group(1)).strip(" ：:，,。;；")
                if value:
                    result[key] = value[:5000]
                    break
    return result


def content_sha256(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


class ReportDataService:
    """将产品主档、合同要素和邮箱净值合并为可审计的报表快照。"""

    def get_product(self, session: Session, product_id: int) -> FundProduct:
        product = session.scalar(
            select(FundProduct).where(
                FundProduct.id == product_id,
                exists(
                    select(FundNav.id).where(
                        FundNav.tenant_id == FundProduct.tenant_id,
                        or_(
                            FundNav.master_product_code == FundProduct.product_code,
                            FundNav.product_code == FundProduct.product_code,
                        ),
                    )
                ),
            )
        )
        if product is None:
            raise AppError("FUND_PRODUCT_NOT_FOUND", "未找到该基金产品", status_code=404)
        return product

    def fields(self, session: Session, product: FundProduct) -> list[dict[str, Any]]:
        effective = product.effective_profile()
        source_meta = product.source_profile_meta or {}
        manual = product.manual_profile or {}
        values: dict[str, Any] = {
            "product_name": product.product_name,
            "product_code": product.product_code,
            **effective,
            "investment_manager": product.investment_manager_info,
            "investment_strategy": product.investment_strategy_info,
        }
        result: list[dict[str, Any]] = []
        for key, definition in REPORT_FIELD_DEFINITIONS.items():
            is_manual = key in manual
            if key == "investment_manager":
                is_manual = product.investment_manager_manual
            elif key == "investment_strategy":
                is_manual = product.investment_strategy_manual
            meta = source_meta.get(key, {}) if isinstance(source_meta.get(key), dict) else {}
            source_type = "manual" if is_manual else meta.get("source_type")
            if not source_type and key in {"product_name", "product_code"}:
                source_type = "email"
            if not source_type and key in {"investment_manager", "investment_strategy"}:
                source_type = "email" if values.get(key) else None
            result.append({
                "key": key,
                "label": definition["label"],
                "group": definition["group"],
                "value": values.get(key),
                "source_type": source_type,
                "source_reference": meta.get("source_reference"),
                "is_manual": is_manual,
                "editable": definition["editable"],
            })
        return result

    def update_field(
        self,
        product: FundProduct,
        *,
        field_key: str,
        value: str | None,
        restore_source: bool,
    ) -> tuple[Any, Any]:
        definition = REPORT_FIELD_DEFINITIONS.get(field_key)
        if not definition or not definition["editable"]:
            raise AppError("REPORT_FIELD_NOT_EDITABLE", "该字段不允许人工覆盖")
        cleaned = value.strip() if value is not None else None
        if field_key == "investment_manager":
            before = product.investment_manager_info
            product.investment_manager_manual = not restore_source
            product.manual_investment_manager_info = None if restore_source else cleaned
            return before, product.investment_manager_info
        if field_key == "investment_strategy":
            before = product.investment_strategy_info
            product.investment_strategy_manual = not restore_source
            product.manual_investment_strategy_info = None if restore_source else cleaned
            return before, product.investment_strategy_info
        profile = dict(product.manual_profile or {})
        before = product.effective_profile().get(field_key)
        if restore_source:
            profile.pop(field_key, None)
        else:
            profile[field_key] = cleaned
        product.manual_profile = profile
        return before, product.effective_profile().get(field_key)

    def apply_contract_fields(
        self,
        product: FundProduct,
        fields: dict[str, str],
        *,
        document_id: int,
        filename: str,
    ) -> None:
        profile = dict(product.source_profile or {})
        metadata = dict(product.source_profile_meta or {})
        for key, value in fields.items():
            if key == "product_code":
                continue
            if key == "product_name":
                # 主档身份仍以托管邮件/备案代码为准，合同名称仅作为来源说明。
                metadata[key] = {
                    "source_type": "contract",
                    "source_reference": filename,
                    "document_id": document_id,
                    "extracted_value": value,
                }
                continue
            if key == "investment_manager":
                product.source_investment_manager_info = value
            elif key == "investment_strategy":
                product.source_investment_strategy_info = value
            else:
                profile[key] = value
            metadata[key] = {
                "source_type": "contract",
                "source_reference": filename,
                "document_id": document_id,
            }
        product.source_profile = profile
        product.source_profile_meta = metadata

    def build_snapshot(
        self,
        session: Session,
        product: FundProduct,
        *,
        report_date: date | None,
        share_product_code: str | None = None,
    ) -> dict[str, Any]:
        series = self._nav_series(
            session, product, report_date=report_date, share_product_code=share_product_code
        )
        if not series:
            raise AppError("REPORT_NAV_EMPTY", "该产品在报告日期前没有可用净值数据")
        actual_date = series[-1][0]
        field_items = self.fields(session, product)
        values = {item["key"]: item["value"] for item in field_items}
        values["report_date"] = actual_date.isoformat()
        values["latest_unit_nav"] = _decimal_text(series[-1][1])
        values["latest_total_nav"] = _decimal_text(series[-1][2])
        metrics = _performance_metrics(series, values.get("inception_date"))
        return {
            "product_id": product.id,
            "product_code": product.product_code,
            "product_name": product.product_name,
            "report_date": actual_date.isoformat(),
            "fields": values,
            "field_provenance": field_items,
            "performance": metrics,
            "nav_series": [
                {
                    "date": item_date.isoformat(),
                    "unit_nav": _decimal_text(unit_nav),
                    "total_nav": _decimal_text(total_nav),
                }
                for item_date, unit_nav, total_nav in series
            ],
        }

    @staticmethod
    def _nav_series(
        session: Session,
        product: FundProduct,
        *,
        report_date: date | None,
        share_product_code: str | None,
    ) -> list[tuple[date, Decimal, Decimal]]:
        conditions = [
            or_(
                FundNav.master_product_code == product.product_code,
                FundNav.product_code == product.product_code,
            ),
            FundNav.unit_nav.is_not(None),
        ]
        if report_date is not None:
            conditions.append(FundNav.nav_date <= report_date)
        if share_product_code:
            conditions.append(FundNav.product_code == share_product_code.strip().upper())
        rows = list(
            session.scalars(
                select(FundNav)
                .where(*conditions)
                .order_by(FundNav.nav_date, FundNav.share_class, FundNav.product_code, FundNav.id)
            )
        )
        grouped: dict[date, list[FundNav]] = defaultdict(list)
        for row in rows:
            grouped[row.nav_date].append(row)
        result: list[tuple[date, Decimal, Decimal]] = []
        for nav_date, date_rows in grouped.items():
            preferred = next(
                (row for row in date_rows if row.share_class == "总份额"),
                date_rows[0],
            )
            if preferred.unit_nav is None:
                continue
            result.append((nav_date, preferred.unit_nav, preferred.total_nav or preferred.unit_nav))
        return result


def _performance_metrics(
    series: list[tuple[date, Decimal, Decimal]], inception_date_value: str | None
) -> dict[str, str | None]:
    dates = [item[0] for item in series]
    values = [float(item[2]) for item in series]
    latest_date = dates[-1]
    latest_value = values[-1]

    def period_return(start_date: date) -> float | None:
        candidates = [
            (day, value)
            for day, value in zip(dates, values, strict=True)
            if day <= start_date
        ]
        start_value = candidates[-1][1] if candidates else values[0]
        return (latest_value / start_value - 1) if start_value else None

    inception_date = _parse_date(inception_date_value) or dates[0]
    elapsed_days = max((latest_date - inception_date).days, 1)
    since_return = latest_value / values[0] - 1 if values[0] else None
    annualized = (
        (latest_value / values[0]) ** (365.25 / elapsed_days) - 1
        if values[0] > 0 and latest_value > 0
        else None
    )
    weekly_values = _weekly_samples(series)
    weekly_returns = [
        weekly_values[index] / weekly_values[index - 1] - 1
        for index in range(1, len(weekly_values))
        if weekly_values[index - 1]
    ]
    volatility = pstdev(weekly_returns) if len(weekly_returns) >= 2 else 0
    sharpe = mean(weekly_returns) / volatility * math.sqrt(52) if volatility else None
    year_start = date(latest_date.year, 1, 1)
    return {
        "annualized_return": _percent(annualized),
        "sharpe_ratio": None if sharpe is None else f"{sharpe:.2f}",
        "return_1m": _percent(period_return(latest_date - timedelta(days=30))),
        "return_3m": _percent(period_return(latest_date - timedelta(days=91))),
        "return_6m": _percent(period_return(latest_date - timedelta(days=183))),
        "return_ytd": _percent(period_return(year_start)),
        "return_1y": _percent(period_return(latest_date - timedelta(days=365))),
        "return_since": _percent(since_return),
        "max_drawdown_ytd": _percent(_max_drawdown(dates, values, year_start)),
        "max_drawdown_since": _percent(_max_drawdown(dates, values, dates[0])),
    }


def _weekly_samples(series: list[tuple[date, Decimal, Decimal]]) -> list[float]:
    buckets: dict[tuple[int, int], float] = {}
    for nav_date, _unit, total in series:
        iso = nav_date.isocalendar()
        buckets[(iso.year, iso.week)] = float(total)
    return list(buckets.values())


def _max_drawdown(dates: list[date], values: list[float], start_date: date) -> float | None:
    selected = [value for day, value in zip(dates, values, strict=True) if day >= start_date]
    if not selected:
        return None
    peak = selected[0]
    drawdown = 0.0
    for value in selected:
        peak = max(peak, value)
        if peak:
            drawdown = min(drawdown, value / peak - 1)
    return drawdown


def _parse_date(value: str | None) -> date | None:
    if not value:
        return None
    normalized = value.replace("年", "-").replace("月", "-").replace("日", "")
    normalized = normalized.replace("/", "-").replace(".", "-")
    try:
        parts = [int(part) for part in normalized.split("-")]
        return date(parts[0], parts[1], parts[2])
    except (ValueError, IndexError):
        return None


def _percent(value: float | None) -> str | None:
    return None if value is None else f"{value * 100:.2f}%"


def _decimal_text(value: Decimal | None) -> str | None:
    return None if value is None else format(value, "f")


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _normalize_contract_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n").replace("\u3000", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
